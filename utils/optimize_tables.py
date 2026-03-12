import docx
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt

def set_cell_border(cell, **kwargs):
    """
    Set cell`s border
    Usage:
    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#000000", "space": "0"},
        bottom={"sz": 12, "color": "#000000", "val": "single"},
        start={"sz": 24, "val": "dashed", "shadow": "true"},
        end={"sz": 12, "val": "dashed"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))

def optimize_tables(docx_path, out_path):
    doc = docx.Document(docx_path)
    
    for i, table in enumerate(doc.tables):
        if i == 0:
            continue # Skip cover page table
        
        # Center the table itself
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Remove table style that might override cell borders
        table.style = 'Normal Table'
        
        for row_idx, row in enumerate(table.rows):
            for cell in row.cells:
                # Center text in cell
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.name = 'Times New Roman'
                        run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
                
                # Clear all borders first by setting them to 'nil' or None
                set_cell_border(
                    cell,
                    top={"val": "nil"},
                    bottom={"val": "nil"},
                    start={"val": "nil"},
                    end={"val": "nil"},
                    insideH={"val": "nil"},
                    insideV={"val": "nil"}
                )
                
                # Top border for the first row (1.5 pt = 12 sz)
                if row_idx == 0:
                    set_cell_border(cell, top={"sz": 12, "val": "single", "color": "000000"})
                    set_cell_border(cell, bottom={"sz": 6, "val": "single", "color": "000000"})
                
                # Bottom border for the last row (1.5 pt = 12 sz)
                if row_idx == len(table.rows) - 1:
                    set_cell_border(cell, bottom={"sz": 12, "val": "single", "color": "000000"})

    doc.save(out_path)
    print(f"Tables optimized and saved to {out_path}")

if __name__ == "__main__":
    in_path = r"D:\02_Study\01_Schoolwork\Graduation Design\docs\midterm\中期 - v2.0_彻底优化.docx"
    out_path = r"D:\02_Study\01_Schoolwork\Graduation Design\docs\midterm\中期 - v2.0_完全优化版.docx"
    optimize_tables(in_path, out_path)
