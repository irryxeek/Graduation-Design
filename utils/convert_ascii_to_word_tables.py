import docx
import re
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc_path = r"D:\02_Study\01_Schoolwork\Graduation Design\docs\midterm\中期 - v2.0.docx"
out_path = r"D:\02_Study\01_Schoolwork\Graduation Design\docs\midterm\中期 - v2.0_彻底优化.docx"

doc = docx.Document(doc_path)

def parse_ascii_table(paragraphs):
    top_line = paragraphs[0].text
    col_starts = [i for i, c in enumerate(top_line) if c in ('┌', '┬')]
    
    rows_data = []
    for p in paragraphs:
        text = p.text
        if not text.strip():
            continue
        if '├' in text or ('─' * 5) in text:
            continue
            
        cells = []
        for i in range(len(col_starts)):
            start = col_starts[i]
            if i + 1 < len(col_starts):
                end = col_starts[i+1]
            else:
                # Find the last boundary character
                end = max(text.rfind('│'), text.rfind('┐'), text.rfind('┘'))
                if end <= start:
                    end = len(text)
            
            # Safe slice
            if start < len(text) and end <= len(text):
                cell_text = text[start+1:end].strip()
            else:
                cell_text = ""
            cells.append(cell_text)
            
        rows_data.append(cells)
        
    cleaned_rows = []
    for row in rows_data:
        if any(cell for cell in row):
            cleaned_rows.append(row)
    return cleaned_rows

blocks_to_replace = []
current_block = []
in_table = False

for p in doc.paragraphs:
    text = p.text.strip()
    if text.startswith('┌'):
        in_table = True
        current_block = [p]
    elif in_table:
        current_block.append(p)
        if text.startswith('└'):
            in_table = False
            blocks_to_replace.append(current_block)
            current_block = []

print(f"Found {len(blocks_to_replace)} ASCII tables.")

for block in blocks_to_replace:
    rows_data = parse_ascii_table(block)
    if not rows_data:
        continue
    
    num_cols = len(rows_data[0])
    
    if num_cols == 1:
        # Flowchart box - merge into one cell
        combined_text = "\n".join([row[0] for row in rows_data if row[0]])
        table = doc.add_table(rows=1, cols=1)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = table.cell(0, 0)
        cell.text = combined_text
        for cp in cell.paragraphs:
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        # Multi-column table
        num_rows = len(rows_data)
        table = doc.add_table(rows=num_rows, cols=num_cols)
        table.style = 'Table Grid' # Standard grid, we can apply better styling later
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        for r_idx, row_data in enumerate(rows_data):
            for c_idx, cell_data in enumerate(row_data):
                if c_idx < len(table.columns):
                    cell = table.cell(r_idx, c_idx)
                    cell.text = cell_data
                    for cp in cell.paragraphs:
                        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Insert before the first paragraph
    first_p = block[0]._p
    first_p.addprevious(table._tbl)
    
    # Remove old paragraphs
    for p in block:
        p_element = p._p
        p_element.getparent().remove(p_element)

doc.save(out_path)
print("Done converting ASCII tables to Word tables.")
