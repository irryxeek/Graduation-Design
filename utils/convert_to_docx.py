import re
import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def create_docx(txt_file, docx_file):
    doc = docx.Document()
    
    # Configure default font for Chinese
    style = doc.styles['Normal']
    font = style.font
    font.name = 'SimSun'
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    font.size = Pt(12)

    with open(txt_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    is_code_block = False
    
    for i, line in enumerate(lines):
        text = line.rstrip('\n')
        
        if not text.strip():
            doc.add_paragraph()
            continue
            
        if text.startswith('====='):
            doc.add_page_break()
            continue
            
        # Title handling
        if i == 0:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            run.font.size = Pt(18)
            run.font.bold = True
            run.font.name = 'SimHei'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
            continue
        elif i == 1:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            run.font.size = Pt(24)
            run.font.bold = True
            run.font.name = 'SimHei'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
            doc.add_paragraph() # Add space
            continue
            
        # Metadata
        if i >= 3 and i <= 9:
            p = doc.add_paragraph(text)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue
            
        # Box drawing characters or fixed width content
        is_table_or_box = any(c in text for c in ['┌', '─', '│', '├', '└', '┐', '┘', '┤', '┬', '┴', '┼', '▼'])
        
        if is_table_or_box:
            is_code_block = True
        elif is_code_block and not text.startswith('  ') and not '|' in text:
            # Maybe the block ended
            is_code_block = False
            
        if is_code_block or is_table_or_box:
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.font.name = 'Consolas'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')
            run.font.size = Pt(10)
            is_code_block = True
            continue

        # Headings
        if re.match(r'^\d+ [^0-9]', text):
            h = doc.add_heading('', level=1)
            run = h.add_run(text)
            run.font.name = 'SimHei'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
            continue
        elif re.match(r'^\d+\.\d+ [^0-9]', text):
            h = doc.add_heading('', level=2)
            run = h.add_run(text)
            run.font.name = 'SimHei'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
            continue
        elif re.match(r'^\d+\.\d+\.\d+ [^0-9]', text):
            h = doc.add_heading('', level=3)
            run = h.add_run(text)
            run.font.name = 'SimHei'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
            continue
            
        # Bullet points
        if text.strip().startswith('·'):
            p = doc.add_paragraph(text.strip()[1:].strip(), style='List Bullet')
            continue
            
        # Normal paragraph
        p = doc.add_paragraph(text)
        
    doc.save(docx_file)

if __name__ == "__main__":
    txt_path = r"D:\02_Study\01_Schoolwork\Graduation Design\docs\midterm\林逸飞-220110814-本科毕业设计中期报告_latex.txt"
    docx_path = r"D:\02_Study\01_Schoolwork\Graduation Design\docs\midterm\林逸飞-220110814-本科毕业设计中期报告_latex.docx"
    create_docx(txt_path, docx_path)
    print("Done")