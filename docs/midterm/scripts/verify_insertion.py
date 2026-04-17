import docx
doc = docx.Document(r'D:\02_Study\01_Schoolwork\Graduation Design\docs\midterm\中期 - v2.0_fixed_final.docx')
p = doc.paragraphs[176] # This is the caption, so image should be 175
p_img = doc.paragraphs[175]
print(f'P175 Text: {p_img.text}')
print(f'P175 Images: {len(p_img._element.xpath(".//w:drawing"))}')
print(f'P176 Text: {p.text}')
