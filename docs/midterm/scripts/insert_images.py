from pathlib import Path
import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

BASE_DIR = Path(__file__).resolve().parents[1]
doc_path = BASE_DIR / 'reports' / '中期 - v2.0_fixed.docx'
doc = docx.Document(str(doc_path))

image_mappings = [
    {
        "target": "数据处理流水线",
        "path": str(BASE_DIR / 'figures' / '数据处理流程图.png'),
        "caption": "图 2.1 FY-3D GNOS L2 ATP 数据预处理流水线",
        "width": 5.5
    },
    {
        "target": "气压对数变换的关键作用",
        "path": str(BASE_DIR / 'figures' / 'pressure_dist_comparison.png'),
        "caption": "图 2.2 气压数据对数变换前后的分布特征对比",
        "width": 5.8
    },
    {
        "target": "训练过程分析",
        "path": str(BASE_DIR / 'figures' / 'loss_curve.png'),
        "caption": "图 2.3 增强型条件 U-Net 扩散模型训练与验证损失收敛曲线",
        "width": 5.0
    },
    {
        "target": "整体评估结果",
        "path": r"D:\02_Study\01_Schoolwork\Graduation Design\evaluation_results_ddim_enhanced\pressure_comparison.png",
        "caption": "图 2.4 测试集气压反演值与真实标签散点对比图",
        "width": 4.5
    },
    {
        "target": "分高度层误差分析",
        "path": r"D:\02_Study\01_Schoolwork\Graduation Design\outputs\figures\Best_RMSE_3.15.png",
        "caption": "图 2.5 典型掩星事件温度与气压反演廓线精细重构示例",
        "width": 4.5
    }
]

def add_image_after_paragraph(para, img_path, caption, width_inches):
    # Create a new paragraph after the current one
    new_p_img = para._element.addnext(docx.oxml.shared.OxmlElement('w:p'))
    new_para_img = docx.text.paragraph.Paragraph(new_p_img, para._parent)
    new_para_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    run = new_para_img.add_run()
    run.add_picture(img_path, width=Inches(width_inches))
    
    # Add caption paragraph
    new_p_cap = new_p_img.addnext(docx.oxml.shared.OxmlElement('w:p'))
    new_para_cap = docx.text.paragraph.Paragraph(new_p_cap, para._parent)
    new_para_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    cap_run = new_para_cap.add_run(caption)
    cap_run.font.size = Pt(10.5) # 五号字
    cap_run.font.name = '宋体'
    rPr = cap_run._element.get_or_add_rPr()
    rFonts = docx.oxml.shared.OxmlElement('w:rFonts')
    rFonts.set(docx.oxml.ns.qn('w:ascii'), 'Times New Roman')
    rFonts.set(docx.oxml.ns.qn('w:eastAsia'), '宋体')
    rPr.append(rFonts)

for mapping in image_mappings:
    if not os.path.exists(mapping["path"]):
        print(f"Warning: Image not found at {mapping['path']}")
        continue
        
    found = False
    for p in doc.paragraphs:
        if mapping["target"] in p.text:
            print(f"Inserting image for: {mapping['target']}")
            add_image_after_paragraph(p, mapping["path"], mapping["caption"], mapping["width"])
            found = True
            break
    if not found:
        print(f"Could not find target paragraph: {mapping['target']}")

doc.save(str(doc_path))
print("Image insertion complete.")
