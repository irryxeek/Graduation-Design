import docx
from docx.oxml.ns import qn
import copy

doc_path = r'D:\02_Study\01_Schoolwork\Graduation Design\docs\midterm\中期 - v2.0_fixed_restored.docx'
doc = docx.Document(doc_path)

if len(doc.tables) < 2:
    print("Error: Reference table not found.")
    exit(1)

ref_table = doc.tables[1]
ref_tblPr = ref_table._element.xpath('w:tblPr')[0]

ascii_data = [
    ["对比维度", "单月模型", "Q1 季度模型", "相对提升"],
    ["训练样本数", "9,244", "26,019", "+181.6%"],
    ["最优验证损失", "0.016882", "0.014280", "-15.4%"],
    ["气压 CC", "0.75", "0.8185", "+9.1%"],
    ["气压 RMSE", "146.82 mb", "123.44 mb", "-15.9%"],
    ["温度 CC", "0.23", "0.2483", "+8.0%"]
]

insert_idx = -1
paras_to_delete = []

in_target_section = False
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if '数据规模对性能的影响' in text:
        in_target_section = True
    elif in_target_section and ('推理验证与性能评估' in text or '2.3.4' in text):
        in_target_section = False
        
    if in_target_section:
        if '┌──' in text or '│ 对比维度' in text or '├──' in text or '└──' in text or '│ 训练样本数' in text or '│ 最优验证损失' in text or '│ 气压 CC' in text or '│ 气压 RMSE' in text or '│ 温度 CC' in text:
            if insert_idx == -1:
                insert_idx = i
            paras_to_delete.append(p)

if insert_idx != -1:
    new_table = doc.add_table(rows=len(ascii_data), cols=len(ascii_data[0]))
    
    # Copy tblPr
    new_tblPr = copy.deepcopy(ref_tblPr)
    new_table._element.replace(new_table._element.xpath('w:tblPr')[0], new_tblPr)
    
    for r_idx, row_data in enumerate(ascii_data):
        row = new_table.rows[r_idx]
        ref_row_idx = 0 if r_idx == 0 else 1
        
        if ref_row_idx < len(ref_table.rows):
            ref_trPr = ref_table.rows[ref_row_idx]._element.xpath('w:trPr')
            if ref_trPr:
                row._element.append(copy.deepcopy(ref_trPr[0]))
        
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = cell_text
            
            if ref_row_idx < len(ref_table.rows) and c_idx < len(ref_table.rows[ref_row_idx].cells):
                ref_col_idx = min(c_idx, len(ref_table.rows[ref_row_idx].cells) - 1)
                ref_cell = ref_table.rows[ref_row_idx].cells[ref_col_idx]
                
                ref_tcPr = ref_cell._element.xpath('w:tcPr')
                if ref_tcPr:
                    if cell._element.xpath('w:tcPr'):
                        cell._element.replace(cell._element.xpath('w:tcPr')[0], copy.deepcopy(ref_tcPr[0]))
                    else:
                        cell._element.insert(0, copy.deepcopy(ref_tcPr[0]))
                
                if ref_cell.paragraphs and cell.paragraphs:
                    ref_p = ref_cell.paragraphs[0]
                    p = cell.paragraphs[0]
                    if ref_p.alignment is not None:
                        p.alignment = ref_p.alignment
                    
                    if ref_p.runs and p.runs:
                        ref_run = ref_p.runs[0]
                        run = p.runs[0]
                        if ref_run.font.name:
                            run.font.name = ref_run.font.name
                            rPr = run._element.get_or_add_rPr()
                            rFonts = rPr.get_or_add_rFonts()
                            rFonts.set(qn('w:eastAsia'), ref_run.font.name)
                        if ref_run.font.size:
                            run.font.size = ref_run.font.size
                        if ref_run.font.bold is not None:
                            run.font.bold = ref_run.font.bold
                        if ref_run.font.color and ref_run.font.color.rgb:
                            run.font.color.rgb = ref_run.font.color.rgb

    p_insert = doc.paragraphs[insert_idx]
    p_insert._p.addprevious(new_table._element)
    
    for p in set(paras_to_delete):
        p_elem = p._element
        parent = p_elem.getparent()
        if parent is not None:
            parent.remove(p_elem)
            
    out_path = r'D:\02_Study\01_Schoolwork\Graduation Design\docs\midterm\中期 - v2.0_fixed.docx'
    doc.save(out_path)
    print("Optimization complete. Target table replaced.")
else:
    print("Could not find ASCII table in the target section.")